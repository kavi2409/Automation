#!/usr/bin/env python3
"""
EFA Proposal Generator - Creates proper Word documents with formatting
Requires: python-docx library
Install: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

def add_border_to_paragraph(paragraph, **kwargs):
    """Add border to paragraph"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '1')
        border.set(qn('w:color'), 'auto')
        pBdr.append(border)
    
    pPr.append(pBdr)

def number_to_words(num):
    """Convert number to words"""
    if num == 0:
        return "Zero"
    
    ones = ['', 'One', 'Two', 'Three', 'Four', 'Five', 'Six', 'Seven', 'Eight', 'Nine']
    tens = ['', '', 'Twenty', 'Thirty', 'Forty', 'Fifty', 'Sixty', 'Seventy', 'Eighty', 'Ninety']
    teens = ['Ten', 'Eleven', 'Twelve', 'Thirteen', 'Fourteen', 'Fifteen', 
             'Sixteen', 'Seventeen', 'Eighteen', 'Nineteen']
    
    def convert_below_thousand(n):
        if n == 0:
            return ''
        elif n < 10:
            return ones[n]
        elif n < 20:
            return teens[n - 10]
        elif n < 100:
            return tens[n // 10] + (' ' + ones[n % 10] if n % 10 != 0 else '')
        else:
            return ones[n // 100] + ' Hundred' + (' and ' + convert_below_thousand(n % 100) if n % 100 != 0 else '')
    
    if num < 1000:
        return convert_below_thousand(num)
    elif num < 1000000:
        thousands = num // 1000
        remainder = num % 1000
        result = convert_below_thousand(thousands) + ' Thousand'
        if remainder > 0:
            result += ' ' + convert_below_thousand(remainder)
        return result
    else:
        millions = num // 1000000
        remainder = num % 1000000
        result = convert_below_thousand(millions) + ' Million'
        if remainder > 0:
            if remainder >= 1000:
                result += ' ' + convert_below_thousand(remainder // 1000) + ' Thousand'
                if remainder % 1000 > 0:
                    result += ' ' + convert_below_thousand(remainder % 1000)
            else:
                result += ' ' + convert_below_thousand(remainder)
        return result

def price_to_words(price):
    """Convert price to words"""
    pounds = int(price)
    pence = round((price - pounds) * 100)
    
    result = number_to_words(pounds) + ' Pounds'
    
    if pence > 0:
        result += ' and ' + number_to_words(pence) + ' Pence'
    
    return result

def get_contract_text(contract_type):
    """Get contract clause text"""
    contracts = {
        '1': 'Our offer is conditional upon the use of the NR26 Professional Service Short Contract (PSSC).',
        '2': 'Our offer is conditional upon the use of the NEC4 Professional Service Short Contract.',
        '3': 'Our offer is conditional upon the use of the NR3 Contract.',
        '4': 'The proposal has been built on the basis that it will be instructed via the PSR:SOW framework and fall under the associated contractual terms. All fees associated with the use of the framework have been incorporated into the prices presented within this proposal.'
    }
    return contracts.get(contract_type, '')

def create_header_table(doc):
    """Create the header table with document info"""
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # Set column widths
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(4.0)
    
    # Row 1
    table.rows[0].cells[0].text = 'Document Ref. No.'
    table.rows[0].cells[1].text = ''
    
    # Row 2
    table.rows[1].cells[0].text = 'Revision'
    table.rows[1].cells[1].text = ''
    
    # Row 3
    table.rows[2].cells[0].text = 'Date'
    table.rows[2].cells[1].text = ''
    
    return table

def create_proposal_document(data):
    """Create the formatted Word document"""
    doc = Document()
    
    # Set up the page
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    
    # Header section with table and "Private and Confidential"
    header_table = create_header_table(doc)
    
    # Add "Private and Confidential" aligned right
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Private and Confidential')
    run.font.size = Pt(10)
    
    doc.add_paragraph()  # Spacing
    
    # FAO section
    p = doc.add_paragraph()
    run = p.add_run(f'FAO: {data["name"]}\n')
    run.font.size = Pt(11)
    run = p.add_run(f'{data["department"]}\n')
    run.font.size = Pt(11)
    run = p.add_run(f'{data["company"]}\n\n')
    run.font.size = Pt(11)
    run = p.add_run(f'{data["date"]}')
    run.font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # "Private and Confidential" again
    p = doc.add_paragraph('Private and Confidential')
    p.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Dear line
    p = doc.add_paragraph(f'Dear {data["name"]},')
    p.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Re: line
    p = doc.add_paragraph()
    run = p.add_run(f'Re: {data["proposal_title"]}')
    run.font.size = Pt(11)
    run.font.bold = True
    
    # Main content box with border
    p = doc.add_paragraph()
    add_border_to_paragraph(p)
    run = p.add_run(f'{data["general_info"]}\n\n')
    run.font.size = Pt(11)
    
    run = p.add_run(f'Our price for the work is £{data["final_price"]:.2f} ({data["price_words"]}). Excluding VAT.\n\n')
    run.font.size = Pt(11)
    
    run = p.add_run(f'{data["pricing_text"]}\n\n')
    run.font.size = Pt(11)
    
    run = p.add_run('Please see full details of scope and deliverables on the following pages:')
    run.font.size = Pt(11)
    
    # Detailed Information box
    p = doc.add_paragraph()
    add_border_to_paragraph(p)
    run = p.add_run(f'{data["detailed_info"]}\n\n')
    run.font.size = Pt(11)
    
    run = p.add_run('Scope\n\n')
    run.font.size = Pt(11)
    run.font.bold = True
    
    run = p.add_run(f'{data["scope"]}\n\n')
    run.font.size = Pt(11)
    run.font.bold = False
    
    run = p.add_run('Deliverables\n\n')
    run.font.size = Pt(11)
    run.font.bold = True
    
    run = p.add_run(f'{data["deliverables"]}\n\n')
    run.font.size = Pt(11)
    run.font.bold = False
    
    run = p.add_run('Resources\n\n')
    run.font.size = Pt(11)
    run.font.bold = True
    
    run = p.add_run(f'{data["resources"]}')
    run.font.size = Pt(11)
    run.font.bold = False
    
    # DURATION section
    p = doc.add_paragraph()
    add_border_to_paragraph(p)
    run = p.add_run('DURATION\n\n')
    run.font.size = Pt(11)
    run.font.bold = True
    
    run = p.add_run(f'The duration of the work is {data["duration"]}, commencing on {data["start_date"]} and concluding on {data["end_date"]}. The project will be billed periodically with the payment application being supported by an up-to-date delivery programme.')
    run.font.size = Pt(11)
    run.font.bold = False
    
    # COMMERCIAL section
    p = doc.add_paragraph()
    add_border_to_paragraph(p)
    run = p.add_run('COMMERCIAL\n\n')
    run.font.size = Pt(11)
    run.font.bold = True
    
    run = p.add_run(f'{data["contract_text"]}\n\n')
    run.font.size = Pt(11)
    run.font.bold = False
    
    run = p.add_run('The ')
    run.font.size = Pt(11)
    run = p.add_run('law of the contract')
    run.font.size = Pt(11)
    run.font.italic = True
    run = p.add_run(' is the Law of England and Wales.\n')
    run.font.size = Pt(11)
    
    run = p.add_run('The ')
    run.font.size = Pt(11)
    run = p.add_run('assessment day')
    run.font.size = Pt(11)
    run.font.italic = True
    run = p.add_run(' is within 28 days from the ')
    run.font.size = Pt(11)
    run = p.add_run('starting date')
    run.font.size = Pt(11)
    run.font.italic = True
    run = p.add_run('.\n')
    run.font.size = Pt(11)
    
    run = p.add_run('The rate for ')
    run.font.size = Pt(11)
    run = p.add_run('delay damages')
    run.font.size = Pt(11)
    run.font.italic = True
    run = p.add_run(' is £0 per day.\n')
    run.font.size = Pt(11)
    
    run = p.add_run('The ')
    run.font.size = Pt(11)
    run = p.add_run('period for reply')
    run.font.size = Pt(11)
    run.font.italic = True
    run = p.add_run(' is 2 weeks.\n\n')
    run.font.size = Pt(11)
    
    run = p.add_run('EFA Engineering holds Professional Indemnity insurance of up to £5 million and Public Liability insurance of up to £5 million. Our liability for any matter is limited to 10% of the contract Price.')
    run.font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Signature section
    p = doc.add_paragraph('Yours sincerely')
    p.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing for signature
    
    p = doc.add_paragraph('Alex Edwards')
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.bold = True
    
    p = doc.add_paragraph('Managing Director')
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.bold = True
    
    p = doc.add_paragraph('+44 (0)7734 646510')
    p.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    p = doc.add_paragraph('EFA Engineering Limited')
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.bold = True
    
    p = doc.add_paragraph('128 City Road,\nLondon,\nUnited Kingdom,\nEC1V 2NX')
    p.runs[0].font.size = Pt(11)
    
    return doc

def get_multiline_input(prompt):
    """Get multiline text input"""
    print(f"\n{prompt}")
    print("(Press Enter twice when done)")
    lines = []
    while True:
        line = input()
        if line == "" and lines and lines[-1] == "":
            lines.pop()
            break
        lines.append(line)
    return "\n".join(lines)

def main():
    os.system('cls' if os.name == 'nt' else 'clear')
    print("=" * 70)
    print(" " * 15 + "EFA PROPOSAL GENERATOR - WORD FORMAT")
    print("=" * 70)
    print()
    
    data = {}
    
    # Basic Information
    print("\n📋 BASIC INFORMATION")
    print("-" * 70)
    data['name'] = input("Name and Surname: ").strip()
    data['department'] = input("Department/Directorate: ").strip()
    data['company'] = input("Company Name: ").strip()
    
    date_input = input("Date (YYYY-MM-DD) [Enter for today]: ").strip()
    data['date'] = date_input if date_input else datetime.date.today().strftime('%Y-%m-%d')
    
    data['proposal_title'] = input("Proposal Title: ").strip()
    
    data['general_info'] = get_multiline_input("General Information and Introduction:")
    
    # Pricing Information
    print("\n💷 PRICING INFORMATION")
    print("-" * 70)
    print("Proposal Type:")
    print("  1. Deliverables")
    print("  2. Timesheets")
    proposal_type = input("Select (1 or 2): ").strip()
    
    data['final_price'] = float(input("Final Price (£): ").strip())
    data['price_words'] = price_to_words(data['final_price'])
    
    pricing_text = ""
    if proposal_type == "1":
        periods = input("Number of weeks/months: ").strip()
        unit = input("Unit (week/month): ").strip()
        rate = data['final_price'] / float(periods)
        pricing_text = f"The price covers a {periods}-{unit} period, based upon £{rate:.2f} per {unit}."
    else:
        print("\nEnter consultant details:")
        consultants = []
        while True:
            job_title = input("  Job Title (or press Enter to finish): ").strip()
            if not job_title:
                break
            charge_rate = input("  Charge Rate (£): ").strip()
            total_shifts = input("  Total Shifts: ").strip()
            consultants.append(f"{job_title} for £{charge_rate} per shift and an anticipated combined number of {total_shifts} total shifts.")
        
        pricing_text = "The price is based upon the below charge out rates and shifts:\n\n" + "\n".join(consultants)
    
    data['pricing_text'] = pricing_text
    
    # Detailed Information
    print("\n📝 DETAILED INFORMATION")
    print("-" * 70)
    data['detailed_info'] = get_multiline_input("Detailed Information:")
    data['scope'] = get_multiline_input("Scope:")
    data['deliverables'] = get_multiline_input("Deliverables:")
    data['resources'] = get_multiline_input("Resources:")
    
    # Duration
    print("\n📅 DURATION")
    print("-" * 70)
    data['duration'] = input("Duration (e.g., '4 weeks' or '6 months'): ").strip()
    data['start_date'] = input("Starting Date (YYYY-MM-DD): ").strip()
    data['end_date'] = input("Ending Date (YYYY-MM-DD): ").strip()
    
    # Commercial Terms
    print("\n📄 COMMERCIAL TERMS")
    print("-" * 70)
    print("Contract Type:")
    print("  1. NR26 (PSSC)")
    print("  2. NEC4 (Professional Service Short Contract)")
    print("  3. NR3 Contract")
    print("  4. PSR:SOW Framework")
    contract_choice = input("Select (1-4): ").strip()
    data['contract_text'] = get_contract_text(contract_choice)
    
    # Generate Document
    print("\n" + "=" * 70)
    print("Generating Word document...")
    print("=" * 70)
    
    doc = create_proposal_document(data)
    
    # Save document
    filename = f"Proposal_{data['proposal_title'].replace(' ', '_')}_{data['date']}.docx"
    doc.save(filename)
    
    print(f"\n✅ SUCCESS! Word document created: {filename}")
    print(f"📁 Location: {os.path.abspath(filename)}")
    print("\n" + "=" * 70)

if __name__ == "__main__":
    try:
        main()
    except ImportError:
        print("\n❌ ERROR: python-docx library not found!")
        print("\nPlease install it by running:")
        print("  pip install python-docx")
        print("\nThen run this script again.")
    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
